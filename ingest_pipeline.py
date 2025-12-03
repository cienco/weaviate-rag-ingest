import os
import io
import base64
from datetime import datetime, timezone
from typing import List, Dict, Any, Optional

import weaviate
from weaviate.classes.init import Auth
from weaviate.classes.config import (
    Property,
    DataType,
    Configure,
)
from weaviate.classes.query import Filter

import fitz  # pymupdf
import docx  # python-docx
import pandas as pd

from google.oauth2 import service_account
from google.auth.transport.requests import Request
from google.cloud import documentai_v1 as documentai

from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from googleapiclient.errors import HttpError
import time


# =============================================================================
# CONFIGURAZIONE DA ENV VAR
# =============================================================================

# Weaviate
WCS_URL = os.getenv("WEAVIATE_URL")
WCS_API_KEY = os.getenv("WEAVIATE_API_KEY")

if not WCS_URL or not WCS_API_KEY:
    raise RuntimeError("WEAVIATE_URL e WEAVIATE_API_KEY devono essere settate nelle env vars.")

# GCP / Vertex / Document AI
GCP_PROJECT_ID = os.getenv("GCP_PROJECT_ID")
if not GCP_PROJECT_ID:
    raise RuntimeError("GCP_PROJECT_ID deve essere settata nelle env vars.")

VERTEX_LOCATION = os.getenv("VERTEX_LOCATION", "us-central1")
DOCAI_PROJECT_ID = os.getenv("DOCAI_PROJECT_ID", GCP_PROJECT_ID)
DOCAI_LOCATION = os.getenv("DOCAI_LOCATION", "eu")
DOCAI_PROCESSOR_ID = os.getenv("DOCAI_PROCESSOR_ID")

if not DOCAI_PROCESSOR_ID:
    raise RuntimeError("DOCAI_PROCESSOR_ID deve essere settata nelle env vars.")

# Scopes per Document AI + Drive
SCOPES = [
    "https://www.googleapis.com/auth/cloud-platform",
    "https://www.googleapis.com/auth/drive.readonly",
]

# Percorso del file JSON della service account
# pattern compatibile con il tuo MCP:
SA_PATH = (
    os.getenv("VERTEX_SA_PATH")
    or os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
    or "/etc/secrets/weaviate-sa.json"
)

if not os.path.exists(SA_PATH):
    raise RuntimeError(f"Service account non trovata: {SA_PATH}")

print(f"[config] Uso service account da: {SA_PATH}")

base_creds = service_account.Credentials.from_service_account_file(
    SA_PATH,
    scopes=SCOPES,
)
base_creds = base_creds.with_scopes(SCOPES)
# NON refresho qui in modo fisso: lo farò dentro get_weaviate_client()

# Sorgente file: cartella locale (potrà essere montata o popolata da un altro sistema)
SOURCE_BASE_DIR = os.getenv("SOURCE_BASE_DIR", "/data/wind_bilance_files")

# Tipi di file
INDEXABLE_TYPES = {"pdf", "docx", "txt", "png", "tif", "xls"}
IGNORED_TYPES = {"zip", "sql", "doc", "msg"}

# Limite di caratteri per il campo text del multimodal embedding
MAX_TEXT_CHARS = 900

# Limite pagine per Document AI non-imageless mode
DOC_PAGES_LIMIT_NON_IMAGELESS = 15


# =============================================================================
# GOOGLE DRIVE CONFIG
# =============================================================================

# ID della cartella root su Google Drive da cui leggere TUTTI i file Wind Bilance
GDRIVE_FOLDER_ID = os.getenv("GDRIVE_FOLDER_ID")
if not GDRIVE_FOLDER_ID:
    raise RuntimeError("GDRIVE_FOLDER_ID non è settata: serve l'ID della cartella su Google Drive.")

_drive_service = None


def get_drive_service():
    """
    Client Google Drive v3 riusando le stesse credenziali (base_creds).
    """
    global _drive_service
    if _drive_service is None:
        _drive_service = build("drive", "v3", credentials=base_creds)
    return _drive_service


# =============================================================================
# HELPER GENERALI
# =============================================================================

def normalize_ext(name: str) -> str:
    """
    Estrae l'estensione in minuscolo senza punto. Es: 'foo.PDF' -> 'pdf'
    """
    import os
    _, ext = os.path.splitext(name)
    return ext.lower().lstrip(".")


def parse_iso(dt_str: Optional[str]) -> Optional[datetime]:
    if not dt_str:
        return None
    try:
        # Gestisce anche ...Z
        if dt_str.endswith("Z"):
            dt_str = dt_str[:-1] + "+00:00"
        return datetime.fromisoformat(dt_str).astimezone(timezone.utc)
    except Exception:
        return None


def now_iso_utc() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def page_img_to_bytes(page_img) -> bytes:
    buf = io.BytesIO()
    page_img.save(buf, format="PNG")
    return buf.getvalue()


def chunk_text(text: str, max_chars: int = MAX_TEXT_CHARS) -> List[str]:
    text = text or ""
    if len(text) <= max_chars:
        return [text] if text else []
    chunks = []
    start = 0
    while start < len(text):
        end = start + max_chars
        chunks.append(text[start:end])
        start = end
    return chunks


# =============================================================================
# CONNESSIONE WEAVIATE + SCHEMA
# =============================================================================

def _build_vertex_headers(token: str) -> Dict[str, str]:
    """
    Crea gli header per far sì che Weaviate possa chiamare text2vec-google
    usando il token OAuth2 ottenuto dalla service account.
    """
    headers: Dict[str, str] = {
        "X-Goog-Vertex-Api-Key": token,
    }
    # opzionale: per billing/quota esplicita
    if GCP_PROJECT_ID:
        headers["X-Goog-User-Project"] = GCP_PROJECT_ID

    return headers


def get_weaviate_client() -> weaviate.WeaviateClient:
    """
    Connessione a Weaviate Cloud usando:
    - API key del cluster (WCS_API_KEY)
    - token OAuth2 di Vertex messo in X-Goog-Vertex-Api-Key

    Questo replica il pattern del tuo MCP:
    il token viene generato dalla service account (sa.json)
    e passato a Weaviate come "API key" per text2vec-google.
    """
    # 1) Refresh del token OAuth2 dalla service account
    global base_creds
    base_creds.refresh(Request())
    vertex_token = base_creds.token

    if not vertex_token:
        raise RuntimeError("Impossibile ottenere il token Vertex dalla service account.")

    # 2) Header HTTP per le chiamate REST del client Weaviate
    headers = _build_vertex_headers(vertex_token)

    # (facoltativo, come nel MCP: popola anche env per eventuali fallback lato client)
    os.environ["GOOGLE_APIKEY"] = vertex_token
    os.environ["PALM_APIKEY"] = vertex_token

    print(f"[vertex-oauth] usando token Vertex (prefix={vertex_token[:10]}...)")

    client = weaviate.connect_to_weaviate_cloud(
        cluster_url=WCS_URL,
        auth_credentials=Auth.api_key(WCS_API_KEY),
        headers=headers,  # 👈 QUI passa il token a Weaviate
    )

    # 3) Metadata gRPC (per sicurezza, come fai nel MCP)
    try:
        conn = getattr(client, "_connection", None)
        if conn is not None:
            meta_list = [
                ("x-goog-vertex-api-key", vertex_token),
            ]
            if GCP_PROJECT_ID:
                meta_list.append(("x-goog-user-project", GCP_PROJECT_ID))

            # prova a settare vari modalità (come nel tuo serve.py)
            try:
                setattr(conn, "grpc_metadata", meta_list)
            except Exception:
                pass
            try:
                setattr(conn, "_grpc_metadata", meta_list)
            except Exception:
                pass
            if hasattr(conn, "set_grpc_metadata"):
                try:
                    conn.set_grpc_metadata(meta_list)
                except Exception:
                    pass
    except Exception as e:
        print(f"[weaviate] warning: non riesco a settare gRPC metadata: {e}")

    return client


def create_schema_if_needed(client: weaviate.WeaviateClient):
    existing_raw = client.collections.list_all()

    existing: set[str] = set()
    for c in existing_raw:
        # caso client nuovo: c è un oggetto con .name
        if hasattr(c, "name"):
            existing.add(c.name)
        # caso client vecchio: c è già una stringa
        elif isinstance(c, str):
            existing.add(c)
        else:
            # fallback per sicurezza
            existing.add(str(c))

    print(f"[schema] Collection esistenti: {existing}")

    # FileIndexStatus (metadati, non vettoriale)
    if "FileIndexStatus" not in existing:
        client.collections.create(
            name="FileIndexStatus",
            properties=[
                Property(
                    name="sourceId",
                    data_type=DataType.TEXT,
                    description="ID sorgente (es: fileId Drive)",
                ),
                Property(name="name", data_type=DataType.TEXT),
                Property(name="path", data_type=DataType.TEXT),
                Property(name="url", data_type=DataType.TEXT),
                Property(name="fileType", data_type=DataType.TEXT),
                Property(name="lastModified", data_type=DataType.TEXT),
                Property(name="indexedAt", data_type=DataType.TEXT),
                Property(name="isDeleted", data_type=DataType.BOOL),
                Property(name="note", data_type=DataType.TEXT),
            ],
            vectorizer_config=Configure.Vectorizer.none(),  # niente embedding su questa collection
        )
        print("[schema] Creata collection FileIndexStatus")

    # WindChunk: SOLO TESTO, text2vec-google
    if "WindChunk" not in existing:
        if not GCP_PROJECT_ID:
            raise RuntimeError("GCP_PROJECT_ID non è settata: serve per text2vec-google.")

        client.collections.create(
            name="WindChunk",
            properties=[
                Property(name="text", data_type=DataType.TEXT),
                Property(name="sourceId", data_type=DataType.TEXT),
                Property(name="fileName", data_type=DataType.TEXT),
                Property(name="fileType", data_type=DataType.TEXT),
                Property(name="pageIndex", data_type=DataType.INT),
                Property(name="chunkIndex", data_type=DataType.INT),
                Property(name="url", data_type=DataType.TEXT),
            ],
            vectorizer_config=Configure.Vectorizer.text2vec_google(
                project_id=GCP_PROJECT_ID,
                # opzionale: puoi specificare il modello, se la tua versione lo supporta:
                # model="textembedding-gecko@003",
            ),
        )
        print("[schema] Creata collection WindChunk (text2vec-google)")


# =============================================================================
# SORGENTE FILE (filesystem locale)
# =============================================================================

class SourceFile:
    def __init__(self, id: str, name: str, path: str, url: str, last_modified: str):
        self.id = id
        self.name = name
        self.path = path
        self.url = url
        self.last_modified = last_modified


def list_source_files() -> List[SourceFile]:
    """
    Legge tutti i file da una cartella di Google Drive (e sotto-cartelle),
    usando GDRIVE_FOLDER_ID come root.
    """
    if not GDRIVE_FOLDER_ID:
        print("[source] GDRIVE_FOLDER_ID non settata: ritorno lista vuota.")
        return []

    service = get_drive_service()
    root_id = GDRIVE_FOLDER_ID
    files: List[SourceFile] = []
    
    print(f"[source] Inizio lettura file da Google Drive (root={root_id})...")

    # BFS sulle cartelle di Drive: (prefix_path, folder_id)
    queue: List[tuple[str, str]] = [("", root_id)]
    folders_processed = 0
    
    # Progress logging basato su tempo (ogni 10 secondi)
    last_progress_log = time.time()
    PROGRESS_LOG_INTERVAL = 10  # secondi

    while queue:
        path_prefix, folder_id = queue.pop(0)
        folders_processed += 1
        
        page_num = 0
        page_token = None
        
        while True:
            page_num += 1
            # --- RETRY SULLA CHIAMATA DRIVE.files().list() ---
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    resp = service.files().list(
                        q=f"'{folder_id}' in parents and trashed = false",
                        fields="nextPageToken, files(id, name, mimeType, modifiedTime, webViewLink)",
                        pageToken=page_token,
                    ).execute()
                    break  # ok, esco dal for
                except HttpError as e:
                    status = e.resp.status if e.resp else None
                    # retry solo su errori 5xx (Internal Error, etc.)
                    if status and 500 <= status < 600 and attempt < max_retries - 1:
                        wait = 2 ** attempt
                        print(f"[source] Drive 5xx (tentativo {attempt+1}/{max_retries}), retry tra {wait}s: {e}")
                        time.sleep(wait)
                        continue
                    # Se non è 5xx o ho esaurito i retry, rilancio
                    print(f"[source] ERRORE permanete da Drive: {e}")
                    raise
            
            # Log quando il retry ha successo (se c'è stato un retry)
            if attempt > 0:
                print(f"[source] Retry riuscito (cartella #{folders_processed})")
            
            items = resp.get("files", [])
            files_in_page = len(items)
            folders_in_page = sum(1 for item in items if item.get("mimeType") == "application/vnd.google-apps.folder")
            files_in_page_count = files_in_page - folders_in_page

            for item in items:
                mime = item.get("mimeType")
                fid = item["id"]
                name = item["name"]

                # Sottocartella → metti in coda
                if mime == "application/vnd.google-apps.folder":
                    new_prefix = f"{path_prefix}{name}/"
                    queue.append((new_prefix, fid))
                    continue

                # File "normale"
                rel_path = f"{path_prefix}{name}"
                last_modified = item.get("modifiedTime")  # ISO 8601
                url = item.get("webViewLink", "")

                files.append(
                    SourceFile(
                        id=fid,
                        name=name,
                        path=rel_path,
                        url=url,
                        last_modified=last_modified,
                    )
                )

            page_token = resp.get("nextPageToken")
            if not page_token:
                break
        
        # Log progress periodico (ogni 10 secondi)
        current_time = time.time()
        if current_time - last_progress_log >= PROGRESS_LOG_INTERVAL:
            print(
                f"[source] Progress: {folders_processed} cartelle processate, "
                f"{len(files)} file trovati, {len(queue)} cartelle in coda"
            )
            last_progress_log = current_time

    print(f"[source] Trovati {len(files)} file in Google Drive (root={root_id}, {folders_processed} cartelle processate)")
    return files


def download_source_file(file_meta: Dict[str, Any]) -> bytes:
    """
    Scarica il file da Google Drive usando sourceId come fileId.
    """
    service = get_drive_service()
    file_id = file_meta["sourceId"]

    request = service.files().get_media(fileId=file_id)
    buf = io.BytesIO()
    downloader = MediaIoBaseDownload(buf, request)

    done = False
    while not done:
        status, done = downloader.next_chunk()
        # Se vuoi puoi loggare: print("Download %d%%" % int(status.progress() * 100))

    return buf.getvalue()


# =============================================================================
# DOCUMENT AI (OCR)
# =============================================================================

_docai_client = None


def get_docai_client():
    global _docai_client
    if _docai_client is not None:
        return _docai_client
    _docai_client = documentai.DocumentProcessorServiceClient(credentials=base_creds)
    return _docai_client


def run_ocr_on_image_bytes(image_bytes: bytes) -> str:
    client = get_docai_client()

    processor_name = client.processor_path(
        DOCAI_PROJECT_ID,
        DOCAI_LOCATION,
        DOCAI_PROCESSOR_ID,
    )

    raw_document = documentai.RawDocument(
        content=image_bytes,
        mime_type="image/png",
    )

    request = documentai.ProcessRequest(
        name=processor_name,
        raw_document=raw_document,
    )

    result = client.process_document(request=request)
    doc = result.document

    text = doc.text or ""
    text = text.replace("\r\n", "\n").strip()
    return text


def split_pdf_bytes(pdf_bytes: bytes, max_pages: int) -> tuple[list[bytes], int]:
    """
    Spezza un PDF (in bytes) in più PDF, ognuno con al massimo max_pages pagine.
    Restituisce (lista_pdf_chunk_bytes, numero_pagine_totale).
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    num_pages = doc.page_count

    chunks: list[bytes] = []

    for start in range(0, num_pages, max_pages):
        end = min(start + max_pages, num_pages)  # end esclusivo
        new_doc = fitz.open()
        # inseriamo le pagine [start, end-1]
        new_doc.insert_pdf(doc, from_page=start, to_page=end - 1)
        
        # Usa save() su BytesIO invece di tobytes("pdf")
        chunk_buffer = io.BytesIO()
        new_doc.save(chunk_buffer, garbage=4, deflate=True)
        chunk_bytes = chunk_buffer.getvalue()
        new_doc.close()
        
        chunks.append(chunk_bytes)
    
    doc.close()

    return chunks, num_pages


def run_ocr_on_pdf_bytes(pdf_bytes: bytes) -> str:
    """
    OCR di un PDF intero usando Document AI in modalità *non-imageless* (quella "ricca").
    - Se il PDF ha <= DOC_PAGES_LIMIT_NON_IMAGELESS pagine -> una sola chiamata
    - Se ha più pagine -> viene spezzato in chunk e chiamato più volte.
    Restituisce il testo concatenato.
    """
    if not (DOCAI_PROJECT_ID and DOCAI_PROCESSOR_ID):
        raise RuntimeError("DOCAI_PROJECT_ID / DOCAI_PROCESSOR_ID non settati")

    client = get_docai_client()

    # Spezza il PDF in chunk di max N pagine
    chunks_bytes, num_pages = split_pdf_bytes(pdf_bytes, DOC_PAGES_LIMIT_NON_IMAGELESS)
    print(
        f"[docai] PDF ha {num_pages} pagine, verrà spezzato in "
        f"{len(chunks_bytes)} chunk da max {DOC_PAGES_LIMIT_NON_IMAGELESS} pagine (non-imageless)."
    )

    full_text_parts: list[str] = []

    for idx, chunk in enumerate(chunks_bytes, start=1):
        name = client.processor_path(DOCAI_PROJECT_ID, DOCAI_LOCATION, DOCAI_PROCESSOR_ID)

        raw_document = documentai.RawDocument(
            content=chunk,
            mime_type="application/pdf",
        )

        # NOTA: niente process_options.enable_image_extraction=False:
        # restiamo in non-imageless mode, così Document AI lavora "full".
        request = documentai.ProcessRequest(
            name=name,
            raw_document=raw_document,
        )

        print(f"[docai] Invio chunk {idx}/{len(chunks_bytes)} a Document AI (non-imageless)...")
        result = client.process_document(request=request)
        document = result.document

        text = (document.text or "").strip()
        print(f"[docai] Chunk {idx}/{len(chunks_bytes)} OCR completato, len={len(text)}")
        if text:
            full_text_parts.append(text)

    full_text = "\n\n".join(full_text_parts).strip()
    print(f"[docai] OCR PDF COMPLETO: len={len(full_text)}")
    return full_text


# =============================================================================
# FILEINDEXSTATUS: SYNC, QUERY, UPDATE
# =============================================================================

def fetch_all_fileindexstatus(
    coll,
    filters: Optional[Any] = None,
    page_size: int = 200,
):
    """
    Recupera TUTTI gli oggetti di FileIndexStatus paginando con offset.
    """
    all_objs = []
    offset = 0

    while True:
        res = coll.query.fetch_objects(
            limit=page_size,
            offset=offset,
            filters=filters,
        )
        if not res.objects:
            break

        all_objs.extend(res.objects)
        offset += page_size

        print(f"[fileindex] Fetched {len(all_objs)} FileIndexStatus objects finora...")

    return all_objs


def sync_source_to_fileindex(client: weaviate.WeaviateClient):
    """
    Sincronizza il contenuto di Google Drive con FileIndexStatus in modo idempotente.

    - Legge TUTTI gli oggetti esistenti da FileIndexStatus (paginando)
    - Li indicizza per sourceId
    - Per ogni file di Drive:
      - se sourceId già presente -> UPDATE
      - se non presente -> INSERT
    - Marca isDeleted=True per gli oggetti che non sono più presenti in Drive
    """
    coll = client.collections.get("FileIndexStatus")

    # 1) Carico TUTTO il pre-esistente
    existing_by_source: dict[str, Any] = {}
    existing_objs = fetch_all_fileindexstatus(coll)

    for obj in existing_objs:
        props = obj.properties or {}
        sid = props.get("sourceId")
        if not sid:
            continue

        if sid in existing_by_source:
            # Se per qualche motivo dovessero esistere duplicati, non ne creiamo altri
            print(f"[sync] WARN: duplicato pre-esistente di sourceId={sid}, uuid={obj.uuid}")
        else:
            existing_by_source[sid] = obj

    print(f"[sync] FileIndexStatus esistenti (unici per sourceId): {len(existing_by_source)}")

    # 2) File da sorgente (Google Drive)
    src_files = list_source_files()
    print(f"[sync] File trovati in Drive: {len(src_files)}")
    
    print(f"[sync] Inizio sincronizzazione: {len(src_files)} file da processare...")
    print(f"[sync] File esistenti in Weaviate: {len(existing_by_source)}")
    print(f"[sync] Stimati nuovi file: ~{len(src_files) - len(existing_by_source)}")

    seen_ids: set[str] = set()
    updates_count = 0
    inserts_count = 0
    start_time = time.time()

    for idx, sf in enumerate(src_files, start=1):
        source_id = sf.id
        seen_ids.add(source_id)

        file_type = normalize_ext(sf.name)

        props = {
            "sourceId":     source_id,
            "name":         sf.name,
            "path":         sf.path,
            "url":          sf.url,
            "fileType":     file_type,
            "lastModified": sf.last_modified,  # stringa ISO da Drive
            "isDeleted":    False,
        }

        if file_type in IGNORED_TYPES:
            props["note"] = f"ignored: {file_type}"
        else:
            props["note"] = ""

        if source_id in existing_by_source:
            # UPDATE idempotente
            try:
                coll.data.update(
                    uuid=existing_by_source[source_id].uuid,
                    properties=props,
                )
                updates_count += 1
            except Exception as e:
                print(f"[sync] ERRORE update per {sf.name} ({source_id}): {e}")
                raise
        else:
            # INSERT solo se non esiste già
            try:
                coll.data.insert(properties=props)
                inserts_count += 1
            except Exception as e:
                print(f"[sync] ERRORE insert per {sf.name} ({source_id}): {e}")
                raise

        # Log più frequente: ogni 10 file invece di 100
        if idx % 10 == 0:
            elapsed = time.time() - start_time
            rate = idx / elapsed if elapsed > 0 else 0
            remaining = (len(src_files) - idx) / rate if rate > 0 else 0
            print(f"[sync] Progress: {idx}/{len(src_files)} file "
                  f"(updates: {updates_count}, inserts: {inserts_count}, "
                  f"rate: {rate:.1f} file/s, ETA: {remaining:.0f}s)")

        # Log ogni 100 file con più dettagli
        if idx % 100 == 0:
            elapsed = time.time() - start_time
            rate = idx / elapsed if elapsed > 0 else 0
            print(f"[sync] Processati {idx}/{len(src_files)} file da Drive "
                  f"(updates: {updates_count}, inserts: {inserts_count}, "
                  f"tempo: {elapsed:.1f}s, rate: {rate:.2f} file/s)")

    elapsed_total = time.time() - start_time
    print(f"[sync] Completata sincronizzazione: {updates_count} updates, {inserts_count} inserts, "
          f"tempo totale: {elapsed_total:.1f}s")

    # 3) (opzionale ma utile) marca come deleted quelli non più visti
    for sid, obj in existing_by_source.items():
        if sid not in seen_ids:
            coll.data.update(
                uuid=obj.uuid,
                properties={"isDeleted": True},
            )

    print(f"[sync] Sync completata: {len(src_files)} file da Drive gestiti.")


def list_files_to_ingest(client: weaviate.WeaviateClient) -> List[Dict[str, Any]]:
    """
    Seleziona i file da ingestare a partire da FileIndexStatus.

    Criteri:
    - isDeleted == False
    - fileType in INDEXABLE_TYPES
    - indexedAt mancante  -> da ingestare
    - OPPURE lastModified > indexedAt -> da re-ingestare
    """
    coll = client.collections.get("FileIndexStatus")

    # Filtra lato Weaviate per isDeleted == False
    where = Filter.by_property("isDeleted").equal(False)

    print("[list_files_to_ingest] Faccio fetch di tutti i FileIndexStatus non deleted...")
    objs = fetch_all_fileindexstatus(coll, filters=where)
    print(f"[list_files_to_ingest] Oggetti FileIndexStatus letti (non deleted): {len(objs)}")

    files: List[Dict[str, Any]] = []
    skipped_not_indexable = 0

    for obj in objs:
        props = obj.properties or {}
        file_type = (props.get("fileType") or "").lower()

        if file_type not in INDEXABLE_TYPES:
            skipped_not_indexable += 1
            continue

        last_mod = parse_iso(props.get("lastModified"))
        indexed_at = parse_iso(props.get("indexedAt"))

        # Mai indicizzato -> da ingestare
        if indexed_at is None:
            files.append(props)
            continue

        # Se non abbiamo lastModified, per sicurezza non ingestiamo di nuovo
        if last_mod is None:
            continue

        # Modificato dopo l'ultimo ingest -> re-ingest
        if last_mod > indexed_at:
            files.append(props)

    print(f"[list_files_to_ingest] File da ingestare (nuovi/modificati): {len(files)}")
    print(f"[list_files_to_ingest] File saltati perché non indicizzabili (tipo non supportato): {skipped_not_indexable}")
    return files


def mark_file_indexed(client: weaviate.WeaviateClient, source_id: str):
    coll = client.collections.get("FileIndexStatus")
    print(f"[mark_file_indexed] Cerco FileIndexStatus per sourceId={source_id}")
    res = coll.query.fetch_objects(
        filters=Filter.by_property("sourceId").equal(source_id),
        limit=1,
    )
    if not res.objects:
        print(f"[mark_file_indexed][WARN] Nessun FileIndexStatus trovato per sourceId={source_id}")
        return

    obj = res.objects[0]
    ts = now_iso_utc()
    coll.data.update(
        uuid=obj.uuid,
        properties={
            "indexedAt": ts,
            "isDeleted": False,
        },
    )
    print(f"[mark_file_indexed] Aggiornato indexedAt={ts} per sourceId={source_id}, uuid={obj.uuid}")


# =============================================================================
# GESTIONE WINDCHUNK
# =============================================================================

def delete_windchunks_for_file(client: weaviate.WeaviateClient, source_id: str):
    coll = client.collections.get("WindChunk")
    where = Filter.by_property("sourceId").equal(source_id)
    coll.data.delete_many(where=where)


def extract_native_text_by_page(pdf_bytes: bytes) -> List[str]:
    texts = []
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page in doc:
        texts.append(page.get_text("text"))
    return texts


# =============================================================================
# INGEST PER TIPO DI FILE
# =============================================================================

def ingest_pdf(client: weaviate.WeaviateClient, file_meta: Dict[str, Any]):
    source_id = file_meta["sourceId"]
    name = file_meta.get("name")
    print(f"[ingest_pdf] Inizio ingest PDF: {name} ({source_id})")

    pdf_bytes = download_source_file(file_meta)
    print(f"[ingest_pdf] PDF scaricato, size={len(pdf_bytes)} bytes")

    # OCR full-document con Document AI
    full_text = run_ocr_on_pdf_bytes(pdf_bytes)
    full_text = (full_text or "").strip()

    if not full_text:
        print(f"[ingest_pdf] Nessun testo OCR per {name}, skip.")
        return

    # Chunking del testo (adatta alla tua funzione di chunking)
    chunks = chunk_text(full_text, max_chars=900)
    print(f"[ingest_pdf] Numero chunk totali per {name}: {len(chunks)}")

    coll = client.collections.get("WindChunk")
    total_chunks = 0

    for ci, chunk in enumerate(chunks):
        props = {
            "sourceId": source_id,
            "fileName": name,
            "fileType": "pdf",
            "pageIndex": -1,  # non abbiamo più info per pagina; puoi lasciarlo -1 o None
            "chunkIndex": ci,
            "text": chunk,
            "url": file_meta.get("url"),
        }

        coll.data.insert(properties=props)
        total_chunks += 1

    print(f"[ingest_pdf] Fine ingest PDF: {name}, total_chunks={total_chunks}")


def ingest_docx(client: weaviate.WeaviateClient, file_meta: Dict[str, Any]):
    source_id = file_meta["sourceId"]
    file_name = file_meta["name"]
    file_url = file_meta["url"]

    print(f"[ingest_docx] Inizio ingest {file_name} ({source_id})")

    file_bytes = download_source_file(file_meta)
    doc_stream = io.BytesIO(file_bytes)
    document = docx.Document(doc_stream)

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    chunks = chunk_text(full_text)
    coll = client.collections.get("WindChunk")

    for idx, chunk in enumerate(chunks):
        if not chunk.strip():
            continue
        props = {
            "sourceId":   source_id,
            "fileName":   file_name,
            "fileType":   "docx",
            "pageIndex":  0,
            "chunkIndex": idx,
            "text":       chunk,
            "url":        file_url,
        }
        coll.data.insert(properties=props)

    print(f"[ingest_docx] Fine ingest {file_name}: {len(chunks)} chunk")


def ingest_txt(client: weaviate.WeaviateClient, file_meta: Dict[str, Any]):
    source_id = file_meta["sourceId"]
    file_name = file_meta["name"]
    file_url = file_meta["url"]

    print(f"[ingest_txt] Inizio ingest {file_name} ({source_id})")

    file_bytes = download_source_file(file_meta)

    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1", errors="ignore")

    chunks = chunk_text(text)
    coll = client.collections.get("WindChunk")

    for idx, chunk in enumerate(chunks):
        if not chunk.strip():
            continue
        props = {
            "sourceId":   source_id,
            "fileName":   file_name,
            "fileType":   "txt",
            "pageIndex":  0,
            "chunkIndex": idx,
            "text":       chunk,
            "url":        file_url,
        }
        coll.data.insert(properties=props)

    print(f"[ingest_txt] Fine ingest {file_name}: {len(chunks)} chunk")


def ingest_image(client: weaviate.WeaviateClient, file_meta: Dict[str, Any]):
    source_id = file_meta["sourceId"]
    file_name = file_meta["name"]
    file_url = file_meta["url"]
    file_type = file_meta.get("fileType", "").lower()

    print(f"[ingest_image] Inizio ingest {file_name} ({source_id})")

    img_bytes = download_source_file(file_meta)

    ocr_text = run_ocr_on_image_bytes(img_bytes)

    coll = client.collections.get("WindChunk")

    props = {
        "sourceId":   source_id,
        "fileName":   file_name,
        "fileType":   file_type,
        "pageIndex":  0,
        "chunkIndex": 0,
        "text":       ocr_text,
        "url":        file_url,
    }
    coll.data.insert(properties=props)

    print(f"[ingest_image] Fine ingest {file_name}")


def ingest_xls(client: weaviate.WeaviateClient, file_meta: Dict[str, Any]):
    source_id = file_meta["sourceId"]
    file_name = file_meta["name"]
    file_url = file_meta["url"]

    print(f"[ingest_xls] Inizio ingest {file_name} ({source_id})")

    file_bytes = download_source_file(file_meta)
    xls_stream = io.BytesIO(file_bytes)

    sheets = pd.read_excel(xls_stream, sheet_name=None)

    coll = client.collections.get("WindChunk")

    chunk_counter = 0

    for sheet_name, df in sheets.items():
        text_repr = df.to_csv(index=False, sep=";", lineterminator="\n")
        text = f"Sheet: {sheet_name}\n{text_repr}"

        chunks = chunk_text(text)

        for idx, chunk in enumerate(chunks):
            if not chunk.strip():
                continue
            props = {
                "sourceId":   source_id,
                "fileName":   file_name,
                "fileType":   "xls",
                "pageIndex":  0,
                "chunkIndex": chunk_counter,
                "text":       chunk,
                "url":        file_url,
            }
            coll.data.insert(properties=props)
            chunk_counter += 1

    print(f"[ingest_xls] Fine ingest {file_name}: {chunk_counter} chunk da {len(sheets)} sheet")


# =============================================================================
# DISPATCH
# =============================================================================

def ingest_single_file(client: weaviate.WeaviateClient, file_meta: Dict[str, Any]):
    source_id = file_meta["sourceId"]
    name = file_meta.get("name")
    file_type = (file_meta.get("fileType") or "").lower()

    print(f"[ingest_single_file] Inizio per sourceId={source_id}, name={name}, type={file_type}")

    try:
        delete_windchunks_for_file(client, source_id)

        if file_type == "pdf":
            print(f"[ingest_single_file] -> ingest_pdf per {name}")
            ingest_pdf(client, file_meta)
        elif file_type == "docx":
            print(f"[ingest_single_file] -> ingest_docx per {name}")
            ingest_docx(client, file_meta)
        elif file_type == "txt":
            print(f"[ingest_single_file] -> ingest_txt per {name}")
            ingest_txt(client, file_meta)
        elif file_type in {"png", "tif"}:
            print(f"[ingest_single_file] -> ingest_image per {name}")
            ingest_image(client, file_meta)
        elif file_type == "xls":
            print(f"[ingest_single_file] -> ingest_xls per {name}")
            ingest_xls(client, file_meta)
        else:
            print(f"[ingest_single_file] Tipo file NON gestito: {file_type}, name={name}")
            return

        print(f"[ingest_single_file] Chiamo mark_file_indexed per sourceId={source_id}")
        mark_file_indexed(client, source_id)
        print(f"[ingest_single_file] Fine ingest_single_file per {name}")
    except Exception as e:
        print(f"[ingest_single_file][ERROR] Errore durante ingest di {name} ({source_id}): {repr(e)}")
        raise


# =============================================================================
# MAIN
# =============================================================================

def main():
    run_id = now_iso_utc()
    print(f"[main] ===== Avvio run ingest {run_id} =====")

    client = get_weaviate_client()
    print("[main] Client Weaviate creato.")

    try:
        print("[main] Creo/verifico schema Weaviate...")
        create_schema_if_needed(client)
        print("[main] Schema ok.")

        print("[main] Sync Drive -> FileIndexStatus...")
        sync_source_to_fileindex(client)
        print("[main] Sync completata.")

        print("[main] Calcolo lista file da ingestare...")
        files = list_files_to_ingest(client)
        print(f"[main] File da ingest (prima del limite): {len(files)}")

        # Limiti da env
        max_files_str = os.getenv("INGEST_MAX_FILES_PER_RUN", "10")
        max_pdfs_str = os.getenv("INGEST_MAX_PDFS_PER_RUN", "3")

        try:
            max_files = int(max_files_str)
        except ValueError:
            max_files = 10

        try:
            max_pdfs = int(max_pdfs_str)
        except ValueError:
            max_pdfs = 3

        print(f"[main] Limiti configurati: max_files={max_files}, max_pdfs={max_pdfs}")

        # Selezione file rispettando i limiti
        selected: List[Dict[str, Any]] = []
        pdf_count = 0

        for fm in files:
            if len(selected) >= max_files:
                break

            ftype = (fm.get("fileType") or "").lower()

            if ftype == "pdf":
                if pdf_count >= max_pdfs:
                    # abbiamo già processato troppi pdf in questo run
                    continue
                pdf_count += 1

            selected.append(fm)

        print(
            f"[main] File da ingest effettivi in questo run: {len(selected)} "
            f"(di cui pdf={pdf_count})"
        )

        # Esecuzione ingest
        for idx, fm in enumerate(selected, start=1):
            name = fm.get("name")
            source_id = fm.get("sourceId")
            ftype = fm.get("fileType")
            print(
                f"[main] >>> [{idx}/{len(selected)}] Inizio ingest file: "
                f"name={name}, type={ftype}, sourceId={source_id}"
            )

            try:
                ingest_single_file(client, fm)
                print(
                    f"[main] <<< [{idx}/{len(selected)}] "
                    f"Ingest COMPLETATA per: {name}"
                )
            except Exception as e:
                print(
                    f"[main][ERROR] Ingest FALLITA per {name} ({source_id}): {repr(e)}"
                )

        print(f"[main] Run ingest {run_id} completato.")
    finally:
        client.close()
        print("[main] Client Weaviate chiuso.")
        print(f"[main] ===== Fine run ingest {run_id} =====")
